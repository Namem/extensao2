# -*- coding: utf-8 -*-
"""Renderiza o Relatorio Final de Extensao II em Word (.docx) editavel."""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage
import relatorio_conteudo as C

OUT = "C:/Users/Rachid/Desktop/NR/Semestre 2026_1/extensao/ceres-diagnostico/docs/extensao/Namem_Rachid_relatorio_final_extensao.docx"

VERDE = RGBColor(0x1A, 0x3A, 0x1A)
DOUR  = RGBColor(0xA9, 0x6E, 0x08)
CINZA = RGBColor(0x2B, 0x2B, 0x2B)
MUTED = RGBColor(0x77, 0x77, 0x77)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEXVERDE = "1A3A1A"; HEXCARD = "2D5A2D"; HEXZEBRA = "F2EEE6"

doc = Document()
# margens e fonte padrao
for sec in doc.sections:
    sec.left_margin = sec.right_margin = Cm(2)
    sec.top_margin = sec.bottom_margin = Cm(2)
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10)
USABLE = Cm(17)


def shade(el, fill):
    pPr = el.get_or_add_pPr() if el.tag.endswith('}p') else el.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def runs_from_md(paragraph, text):
    """Adiciona runs interpretando **negrito**."""
    for i, part in enumerate(re.split(r'\*\*(.+?)\*\*', text)):
        if part == '':
            continue
        r = paragraph.add_run(part)
        if i % 2 == 1:
            r.bold = True


def add_h1(text):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    shade(p._p, HEXVERDE)
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = WHITE
    # padding via indentacao leve
    p.paragraph_format.left_indent = Pt(4)


def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = VERDE


def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = DOUR


def add_p(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5); p.paragraph_format.line_spacing = 1.12
    runs_from_md(p, text)
    for r in p.runs:
        r.font.color.rgb = CINZA


def add_bullets(items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        runs_from_md(p, it)
        for r in p.runs:
            r.font.size = Pt(10); r.font.color.rgb = CINZA


def set_col_widths(table, fracs):
    for row in table.rows:
        for c, frac in zip(row.cells, fracs):
            c.width = int(USABLE * frac)


def add_table(headers, rows):
    ncol = len(headers)
    t = doc.add_table(rows=1, cols=ncol)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    t.autofit = False
    # header
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        shade(cell._tc, HEXCARD)
        para = cell.paragraphs[0]; para.paragraph_format.space_after = Pt(1)
        r = para.add_run(h); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE
    # body
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            if ri % 2 == 1:
                shade(cells[j]._tc, HEXZEBRA)
            para = cells[j].paragraphs[0]; para.paragraph_format.space_after = Pt(1)
            runs_from_md(para, val)
            for rr in para.runs:
                rr.font.size = Pt(9); rr.font.color.rgb = CINZA
    if ncol == 2:
        fr = [0.32, 0.68] if len(headers[0]) < 14 else [0.5, 0.5]
    elif ncol == 5:
        fr = [0.08, 0.42, 0.15, 0.12, 0.23]
    else:
        fr = [1.0/ncol]*ncol
    set_col_widths(t, fr)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(path, caption):
    w, h = PILImage.open(path).size
    iw = Cm(15.5); ih = iw*h/w
    if ih > Cm(10.5):
        ih = Cm(10.5); iw = ih*w/h
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=iw)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(re.sub(r'\*\*(.+?)\*\*', r'\1', caption)); r.italic = True
    r.font.size = Pt(8.5); r.font.color.rgb = MUTED


def add_code(text):
    p = doc.add_paragraph()
    shade(p._p, "F2EEE6")
    pf = p.paragraph_format; pf.space_before = Pt(2); pf.space_after = Pt(6); pf.left_indent = Pt(6)
    # borda leve
    pPr = p._p.get_or_add_pPr(); bdr = OxmlElement('w:pBdr')
    for side in ('top', 'bottom', 'left', 'right'):
        e = OxmlElement('w:' + side); e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:color'), 'D8D0C0'); e.set(qn('w:space'), '3'); bdr.append(e)
    pPr.append(bdr)
    for i, ln in enumerate(text.split('\n')):
        r = p.add_run(('' if i == 0 else '\n') + ln)
        r.font.name = 'Consolas'; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x1A, 0x2D, 0x1A)


def add_figrow(cols, items):
    t = doc.add_table(rows=0, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    colw = int(USABLE / cols)
    for i in range(0, len(items), cols):
        chunk = items[i:i+cols]
        cells = t.add_row().cells
        for j in range(cols):
            cell = cells[j]; cell.width = colw
            para = cell.paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if j < len(chunk):
                path, cap = chunk[j]
                w, h = PILImage.open(path).size
                iw = colw - Cm(0.3); ih = iw * h / w
                if ih > Cm(8.2):
                    ih = Cm(8.2); iw = ih * w / h
                para.add_run().add_picture(path, width=iw)
                capp = cell.add_paragraph(); capp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                capp.paragraph_format.space_after = Pt(4)
                rr = capp.add_run(re.sub(r'\*\*(.+?)\*\*', r'\1', cap))
                rr.italic = True; rr.font.size = Pt(8); rr.font.color.rgb = MUTED
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── Cabecalho ──
tt = doc.add_paragraph(); tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tt.add_run(C.TITULO); r.bold = True; r.font.size = Pt(19); r.font.color.rgb = VERDE
st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run(C.SUBTITULO); r.font.size = Pt(9.5); r.font.color.rgb = MUTED
hr = doc.add_paragraph(); hr.paragraph_format.space_after = Pt(6)
pbdr = hr._p.get_or_add_pPr(); bdr = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '12'); bot.set(qn('w:color'), HEXVERDE)
bot.set(qn('w:space'), '1'); bdr.append(bot); pbdr.append(bdr)

for blk in C.BLOCKS:
    k = blk[0]
    if k == 'h1': add_h1(blk[1])
    elif k == 'h2': add_h2(blk[1])
    elif k == 'h3': add_h3(blk[1])
    elif k == 'p': add_p(blk[1])
    elif k == 'bul': add_bullets(blk[1])
    elif k == 'tab': add_table(blk[1], blk[2])
    elif k == 'fig': add_figure(blk[1], blk[2])
    elif k == 'code': add_code(blk[1])
    elif k == 'figrow': add_figrow(blk[1], blk[2])

doc.save(OUT)
print("OK:", OUT)
